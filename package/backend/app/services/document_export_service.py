from __future__ import annotations

import io
import json
import os
from dataclasses import dataclass
from typing import List

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from app.config import settings
from app.models.models import OptimizationSegment, OptimizationSession
from app.services.docx_preserve_service import export_preserved_docx

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PDF_MIME = "application/pdf"
TXT_MIME = "text/plain; charset=utf-8"
MD_MIME = "text/markdown; charset=utf-8"


@dataclass
class ExportResult:
    content: bytes
    filename: str
    media_type: str


def get_final_text(segments: List[OptimizationSegment]) -> str:
    return "\n\n".join(
        segment.enhanced_text or segment.polished_text or segment.original_text or ""
        for segment in segments
    )


def build_export(session: OptimizationSession, segments: List[OptimizationSegment], export_format: str) -> ExportResult:
    final_text = get_final_text(segments)
    base_name = _base_filename(session)

    if export_format == "txt":
        return ExportResult(
            content=final_text.encode("utf-8"),
            filename=f"{base_name}.txt",
            media_type=TXT_MIME,
        )

    if export_format == "md":
        if session.source_type == "md" and session.source_file_blob and session.preserve_format_available:
            try:
                final_text = _build_preserved_markdown(session, segments)
            except Exception as exc:
                print(f"[WARN] Markdown preserved export failed, falling back to plain markdown: {exc}", flush=True)
        return ExportResult(
            content=final_text.encode("utf-8"),
            filename=f"{base_name}.md",
            media_type=MD_MIME,
        )

    if export_format == "docx":
        content = _build_docx(session, segments, final_text)
        return ExportResult(
            content=content,
            filename=f"{base_name}.docx",
            media_type=DOCX_MIME,
        )

    if export_format == "pdf":
        return ExportResult(
            content=_build_pdf(final_text),
            filename=f"{base_name}.pdf",
            media_type=PDF_MIME,
        )

    raise ValueError("不支持的导出格式")


def _base_filename(session: OptimizationSession) -> str:
    if session.source_filename:
        stem = os.path.splitext(session.source_filename)[0].strip()
        if stem:
            return f"{stem}_optimized"
    return f"optimized_{session.session_id}"


def _build_docx(session: OptimizationSession, segments: List[OptimizationSegment], final_text: str) -> bytes:
    if session.source_type == "docx" and session.source_file_blob and session.preserve_format_available:
        try:
            return export_preserved_docx(session, segments)
        except Exception as exc:
            print(f"[WARN] Word 保格式导出失败，改用普通 docx 导出: {exc}", flush=True)

    return _build_plain_docx(final_text)


def _build_plain_docx(text: str) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    paragraphs = text.split("\n\n") if text else [""]
    for paragraph_text in paragraphs:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(paragraph_text)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(12)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _final_segment_text(segment: OptimizationSegment) -> str:
    return segment.enhanced_text or segment.polished_text or segment.original_text or ""


def _build_preserved_markdown(session: OptimizationSession, segments: List[OptimizationSegment]) -> str:
    source = session.source_file_blob.decode("utf-8")
    replacements = []

    for segment in segments:
        if not segment.structure_meta:
            continue
        try:
            meta = json.loads(segment.structure_meta)
        except json.JSONDecodeError:
            continue

        start = meta.get("source_offset_start")
        end = meta.get("source_offset_end")
        if not isinstance(start, int) or not isinstance(end, int) or start < 0 or end < start:
            continue
        replacements.append((start, end, _final_segment_text(segment)))

    if not replacements:
        raise ValueError("No Markdown blocks are available for preserved export")

    result_parts = []
    cursor = 0
    for start, end, replacement in sorted(replacements, key=lambda item: item[0]):
        if start < cursor:
            continue
        result_parts.append(source[cursor:start])
        original_block = source[start:end]
        leading = original_block[:len(original_block) - len(original_block.lstrip())]
        trailing = "\n" if original_block.endswith("\n") else ""
        result_parts.append(f"{leading}{replacement}{trailing}")
        cursor = end
    result_parts.append(source[cursor:])
    return "".join(result_parts)


def _build_pdf(text: str) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except ImportError as exc:
        raise ValueError("PDF 导出依赖 reportlab 未安装") from exc

    font_name = _register_pdf_font(pdfmetrics, TTFont, UnicodeCIDFont)
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=42, leftMargin=42, topMargin=42, bottomMargin=42)
    styles = getSampleStyleSheet()
    normal = styles["Normal"]
    normal.fontName = font_name
    normal.fontSize = 11
    normal.leading = 18

    story = []
    for paragraph in (text.split("\n\n") if text else [""]):
        escaped = paragraph.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")
        story.append(Paragraph(escaped or " ", normal))
        story.append(Spacer(1, 8))

    doc.build(story)
    return buffer.getvalue()


def _register_pdf_font(pdfmetrics, TTFont, UnicodeCIDFont) -> str:
    configured_path = getattr(settings, "PDF_FONT_PATH", None)
    configured_name = getattr(settings, "PDF_FONT_NAME", None) or "DocumentChineseFont"
    candidates = [
        configured_path,
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simsun.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
    ]

    for path in candidates:
        if path and os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont(configured_name, path))
                return configured_name
            except Exception:
                continue

    try:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        return "STSong-Light"
    except Exception:
        return "Helvetica"
